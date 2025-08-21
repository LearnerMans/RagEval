from __future__ import annotations

from .registry import register_extension


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document  # python-docx
        import io

        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return ""


register_extension(".docx", _extract_docx)


